import sys
from django.core.management.base import BaseCommand
from docx import Document
import pandas as pd
from openpyxl import load_workbook
from ...models import (Songs,
                       LiturgyBook,
                       LiturgyDetails)
from ...models import (KaraitesBookAsArray,
                       KaraitesBookDetails)
from bs4 import BeautifulSoup
from django.core.files import File
from pathlib import Path

path = Path() / 'data_karaites/Word Documents/Liturgy/'
output_path = Path() / 'data_karaites/Out_xls/Liturgy/'


class Command(BaseCommand):
    """ Export LiturgyBooks in karaites model to xls file that should be updated
        and imported to database by LiturgyBooks.py
    """

    @staticmethod
    def extract_intro(file_name):
        doc = Document(file_name)
        introduction = []

        for element in doc.element.body:
            if element.tag.endswith('p'):
                introduction.append(doc.paragraphs[len(introduction)].text)
            elif element.tag.endswith('tbl'):
                break
        return introduction

    @staticmethod
    def read_docx_table(file_name, num_cols, start_col):

        doc = Document(file_name)
        data = []

        for table in doc.tables:
            # Initialize temp_row to None
            temp_row = ['', '', '']
            # Hebrew is second column
            col = 1
            translation_col = 2
            for i, row in enumerate(table.rows):

                # even row
                if i % 2 == 0:
                    for cell in row.cells:
                        if cell.text:
                            temp_row[col] = cell.text
                            col -= 1
                # odd row
                else:
                    for cell in row.cells:
                        if cell.text:
                            temp_row[translation_col] = cell.text
                    data.append(temp_row)
                    temp_row = ['', '', '']
                    col = 1

        # expand to num_cols and separate hebrew and english in sentences
        line_number = 1
        expanded_data = []
        for row in data:

            hebrew_sentences = row[0].split('\n')
            transliteration_sentences = row[1].split('\n')
            translation_sentences = row[2].split('\n')

            for hebrew, transliteration, english in zip(hebrew_sentences, transliteration_sentences,
                                                        translation_sentences):
                line = [''] * num_cols
                line[8] = line_number
                line[9] = hebrew
                line[10] = transliteration
                line[11] = english
                expanded_data.append(line)
                line_number += 1

        return expanded_data

    def get_book_details(self, file_name):
        for book in KaraitesBookAsArray.objects.filter(book__first_level__first_level='Liturgy'):
            # details
            english_name = book.book.book_title_en
            hebrew_name = book.book.book_title_he
            author = book.book.author
            intro = book.book.introduction
            toc = book.book.toc
            language = book.book.book_language

    def handle(self, *args, **options):
        """ """
        labels = ['File Name', 'Occasion', 'Hebrew Name', 'English Name', 'Display', 'Divisions',
                  'Censored', 'Reciter', 'Hebrew Line #', 'Hebrew Text', 'English Transliteration',
                  'English Translation', 'Comments', 'Musical Display Notes', 'Time Starting', 'Time Ending']

        i = 1
        for file in path.glob('**/*.docx'):

            file_parts = file.parts
            # skip introduction files
            if any('introduction' in s.lower() for s in file_parts):
                continue
            # introduction = self.extract_intro(file)

            out_filename = file_parts[-1].replace('.docx', '.xlsx')

            data = self.read_docx_table(file, len(labels), start_col=8)

            Path(output_path).mkdir(parents=True, exist_ok=True)

            df = pd.DataFrame(data, columns=labels)
            df.to_excel(Path(output_path) / f'{out_filename}', index=False)

            # # add the introduction to another tab
            # with pd.ExcelWriter(Path(output_path) / f'{out_filename}', engine='openpyxl',mode='a') as writer:
            #     df = pd.DataFrame(introduction, columns=['Introduction'])
            #     df.to_excel(writer, sheet_name='Introduction', index=False)

            print(f'{i} - {out_filename} created')
            i += 1
